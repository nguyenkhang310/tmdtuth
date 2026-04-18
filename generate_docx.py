from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def main():
    doc = Document()
    
    # Setup Title
    title = doc.add_heading('BÁO CÁO ĐỒ ÁN / LUẬN VĂN', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph('TỔNG QUAN HỆ THỐNG THƯƠNG MẠI ĐIỆN TỬ UTH STORE')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    # Introduction
    doc.add_heading('LỜI NÓI ĐẦU', level=1)
    doc.add_paragraph('Tài liệu báo cáo này trình bày toàn bộ phương pháp, cấu trúc kỹ thuật và kiến trúc triển khai của hệ thống Thương Mại Điện Tử UTH Store. Báo cáo này được cấu trúc cặn kẽ để dễ dàng trở thành tư liệu mẫu cao cấp cho công tác chấm báo cáo đồ án, phân tích phần mềm và luận văn.')
    
    # Chapter 1
    doc.add_heading('CHƯƠNG 1: TỔNG QUAN DỰ ÁN VÀ NGHIỆP VỤ HỆ THỐNG', level=1)
    doc.add_heading('1.1. Mục Tiêu Dự Án', level=2)
    doc.add_paragraph('UTH Store là nền tảng thương mại trực tuyến hiện đại, ứng dụng chuyên sâu cho việc phân phối ngành hàng thời trang. Dự án đi theo mô hình vòng đời phần mềm tiêu chuẩn với các tác vụ then chốt phải được giải quyết bao gồm:')
    doc.add_paragraph('Tạo quy trình chọn, lọc và duyệt sản phẩm trơn tru.', style='List Bullet')
    doc.add_paragraph('Triển khai logic Giỏ hàng thông minh (Dạng Guest qua Session và User lưu đồng bộ Database).', style='List Bullet')
    doc.add_paragraph('Chuyển đổi số thanh toán: Ví MoMo, VNPAY, COD và Bank Transfer.', style='List Bullet')
    
    doc.add_heading('1.2. Giải pháp Công nghệ (Tech Stack)', level=2)
    doc.add_paragraph('Hệ thống được phát triển chuyên sâu với ngôn ngữ Python (Flask Framework) áp dụng cơ chế thiết kế MVC tiêu chuẩn hiện đại.')
    t = doc.add_table(rows=1, cols=2)
    t.style = 'Table Grid'
    hdr_cells = t.rows[0].cells
    hdr_cells[0].text = 'Phân Hệ'
    hdr_cells[1].text = 'Công Nghệ Sử Dụng'
    data = [
        ('Backend (Core)', 'Python 3, Flask'),
        ('Lớp Database', 'SQLite, SQLAlchemy ORM, Flask-Migrate'),
        ('Bảo Mật & RateLimit', 'Werkzeug, Flask-Login, Flask-WTF, Flask-Limiter, Redis'),
        ('Giao diện (Frontend)', 'Jinja2, HTML5, Bootstrap, jQuery JS, AJAX API')
    ]
    for ph, cn in data:
        row_cells = t.add_row().cells
        row_cells[0].text = ph
        row_cells[1].text = cn
    
    doc.add_page_break()

    # Chapter 2
    doc.add_heading('CHƯƠNG 2: PHÂN TÍCH KIẾN TRÚC VÀ CƠ SỞ DỮ LIỆU', level=1)
    doc.add_heading('2.1. Cấu trúc thực thể đối tượng (ER - Entity Relationship)', level=2)
    doc.add_paragraph('Khối dữ liệu của nền tảng được trừu tượng hóa qua cơ chế ORM của SQLAlchemy rất rõ ràng để truy vết và tái cấu trúc sang MySQL hay Postgres tùy chọn:')
    doc.add_paragraph('User: Nắm giữ định danh (username, password băm chuẩn thuật toán PBKDF2), role phân quyền nhánh admin/user.', style='List Number')
    doc.add_paragraph('Product & Category: Cấu trúc cha con 1-N. Đặc thù sản phẩm được kiểm soát qua trường Tồn Kho (Stock) ngăn ngừa thuật toán Over-selling trong bán hàng.', style='List Number')
    doc.add_paragraph('Order, OrderItem & PaymentTransaction: Ghi log giao dịch đa tầng giúp dò vết toàn diện. Ghi nhận snapshot hệ giá tiền gốc và lưu trữ phản hồi chuỗi JSON (raw response) chống chối bỏ hóa đơn (Momo, VNPay).', style='List Number')
    doc.add_paragraph('Cart & CartItem: Cơ chế Giỏ hàng liên kết master-detail hỗ trợ đồng bộ hóa từ máy trạm di động qua thiết bị Desktop.', style='List Number')

    doc.add_heading('2.2. Kiểm soát luồng hoạt động (Request-Response Lifecycle)', level=2)
    doc.add_paragraph('Bộ điều phối Web nhận Request -> Chuyển qua lớp chặn thư rác của Limiter(Redis) -> Vượt qua xác thực CSRF Token -> Xử lý logic tại hàm của file app.py -> Trộn View bằng Engine Templating Jinja2 -> HTML mã hóa về trình duyệt người dùng.')

    doc.add_page_break()

    # Chapter 3
    doc.add_heading('CHƯƠNG 3: PHÂN TÍCH CHUYÊN SÂU SOURCE CODE', level=1)
    
    doc.add_heading('3.1. Kernel Điều Tuyến Lõi - app.py', level=2)
    doc.add_paragraph('File gốc duy trì hơn 1000 dòng lệnh để vận hành tất cả endpoint. Chức năng chính thống gồm:')
    doc.add_paragraph('Cơ chế gộp giỏ hàng (Cart Merge System): _merge_session_cart_into_user_cart() - Phân giải trạng thái cực mạnh. Guest lưu mã Session, khi định danh thanh toán thành Register User, hàm sẽ đồng nhất hóa mảng Session sang DB không gây lặp đè.', style='List Bullet')
    doc.add_paragraph('Cơ chế Mã Khuyến Mãi (Coupon Engine): Tính toán đa luồng từ cắt Tỉ lệ phần trăm, trừ mức giá Fix và chiết khấu Free Ship nội đô thành một tổng tuyến duy nhất.', style='List Bullet')
    doc.add_paragraph('Hệ thống HMAC Security: Khóa và tạo mã kiểm dò thông báo chuyển khoản API webhook để chống gian lận thay đổi dữ liệu.', style='List Bullet')

    doc.add_heading('3.2. Mapping Database - models.py & extensions.py', level=2)
    doc.add_paragraph('Các đối tượng CSDL quy chuẩn. Hàm tham chiếu Backref trong Python cho phép việc Left/Right Join giữa các table hoàn toàn ngầm không bắt coder thao tác SQL dễ phát sinh lỗi.', style='List Bullet')
    doc.add_paragraph('Thư mục models được tách biệt Singleton bằng file modules extensions tránh lỗi Vòng lặp thư viện (Circular Import Dependency).', style='List Bullet')

    doc.add_heading('3.3. Layer Sinh Dữ Liệu Ảo Mock - seed.py', level=2)
    doc.add_paragraph('Chuẩn hóa tiến trình làm đầy CSDL trắng thành mạng lưới hàng mẫu và quản trị viên, giúp bảo vệ môi trường dev và quy chuẩn hóa Unit Testing.')

    doc.add_page_break()

    # Chapter 4
    doc.add_heading('CHƯƠNG 4: PHÂN HỆ GIAO DIỆN (TEMPLATES DIRECOTRY)', level=1)
    doc.add_paragraph('Không thiết kế rời rạc dạng HTML thường, Jinja2 được dùng định hình bộ giao diện Components quy mô lớn:')
    doc.add_paragraph('Framework Khung (base.html): Điểm cài thư viện Front-End, điều phối Flashing Alerts. Các file HTML con chỉ cần extend từ file gốc.', style='List Bullet')
    doc.add_paragraph('Tương tác khách (products.html, detail.html): Liên túc dùng AJAX chặn quá tải và hiển thị thanh số lượng có validation chống ép kiểu (Type Casting Exploit).', style='List Bullet')
    doc.add_paragraph('Module Tín Dụng (checkout.html): Logic ẩn hiện module điền thông tin khi qua bank hay ví online.', style='List Bullet')
    doc.add_paragraph('Control Panel Admin (admin.html, my-orders.html): Trích xuất và tô màu trực quan (Color Badge) trạng thái của Order để dễ phục vụ Vận Đơn.', style='List Bullet')

    doc.add_page_break()

    # Chapter 5
    doc.add_heading('CHƯƠNG 5: KIẾN TRÚC BẢO MẬT VÀ TỐI ƯU HIỆU SUẤT TRÊN RỘNG', level=1)
    doc.add_heading('5.1. Khía Cạnh Bảo Mật An Toàn Mạng', level=2)
    doc.add_paragraph('Dự án chú trọng không gian an toàn thương mại:')
    doc.add_paragraph('Băm mật khẩu chuỗi dài qua cơ chế cao cấp Werkzeug (PBKDF2 HMAC SHA-256). Kể cả bị mất file DB cũng không khai thác ngược chuỗi mật khẩu.', style='List Number')
    doc.add_paragraph('Chống chèn mã rác (Injection/XSS) qua thư viện file tĩnh và String sanitizer.', style='List Number')
    doc.add_paragraph('Tokens Anti-CSRF trên từng trang điền Form, phá nát các kỹ thuật link fishing giả lập request.', style='List Number')

    doc.add_heading('5.2. Chống DDOS bằng bộ nhớ đệm Limit On-RAM', level=2)
    doc.add_paragraph('Việc tích hợp Flask-Limiter thông qua in-memory cache của hệ quản trị Redis triệt tiêu thời gian phản ứng HDD. Biện pháp này bẻ gãy mọi cố ý DDoS hoặc spam đăng nhập với hạn mức 200 lượt/IP/giờ.')

    # Conclusion
    doc.add_heading('KẾT LUẬN VÀ HƯỚNG PHÁT TRIỂN', level=1)
    doc.add_paragraph('Ứng dụng đã hoàn thành tính năng của hệ sinh thái thương mại trực tuyến, minh chứng cho việc mô hình hóa cấu trúc MVC bền vững trên nền tảng Python Flask.')
    doc.add_paragraph('Hướng phát triển tiếp theo có thể đi lên theo mô hình hóa Micro-Service bằng cách tích hợp trực tiếp qua API RESTful và bọc (Containerized) hệ thống bằng Docker để triển khai Load-balancer Kubernetes.')

    doc.save('Bao_Cao_Tong_Quan_UTH_Store.docx')
    print("SUCCESS: File .docx has been generated.")

if __name__ == '__main__':
    main()
